# For reading file name
import glob
print(glob.glob('.'))
myList=[]
myListWithoutSlash=[]
for file_name in glob.iglob('./Imag/*.*', recursive=True):
  print(file_name)
  myList.append(file_name)
  # using split() 
  # Get String after substring occurrence 
  res = file_name.split('/')
  for item in res:
      if "jpeg" in item:
        myListWithoutSlash.append(item)
      if "JPEG" in item:
        myListWithoutSlash.append(item)
      if "png" in item:
        myListWithoutSlash.append(item)
      if "PNG" in item:
        myListWithoutSlash.append(item)
      if "jpg" in item:
        myListWithoutSlash.append(item)
      if "JPG" in item:
        myListWithoutSlash.append(item)
# for python-docx package
# create words documents
from docx import Document
from docx.shared import Cm
document = Document()
table = document.add_table(rows = 1, cols = 2)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'ImageName'
hdr_cells[1].text = 'Image'
for i in range(len(myList)): 
  print(i)
  print(myListWithoutSlash[i])
  row_cells = table.add_row().cells
  row_cells[0].text = myListWithoutSlash[i]
  p = row_cells[1].add_paragraph()
  r = p.add_run()
  r.add_picture(myList[i],width=Cm(4.0), height=Cm(4))
document.save('./demo.docx')


# # For reading file name
# import glob
# print(glob.glob('.'))
# myList=[]
# myListWithoutSlash=[]
# for file_name in glob.iglob('./*.*', recursive=True):
#   print(file_name)
#   myList.append(file_name)
#   # using split() 
#   # Get String after substring occurrence 
#   res = file_name.split('/')
#   for item in res:
#       if "jpeg" in item:
#         myListWithoutSlash.append(item)
#       if "JPEG" in item:
#         myListWithoutSlash.append(item)
#       if "png" in item:
#         myListWithoutSlash.append(item)
#       if "PNG" in item:
#         myListWithoutSlash.append(item)
#       if "jpg" in item:
#         myListWithoutSlash.append(item)
#       if "JPG" in item:
#         myListWithoutSlash.append(item)
# # for python-docx package
# # create words documents
# from docx import Document
# from docx.shared import Cm
# document = Document()
# table = document.add_table(rows = 1, cols = 4)
# table.style = 'Table Grid'
# hdr_cells = table.rows[0].cells
# hdr_cells[0].text = 'ProductName'
# hdr_cells[1].text = 'Id'
# hdr_cells[2].text = 'ImageName'
# hdr_cells[3].text = 'Image'
# for i in range(len(myList)): 
#   print(i)
#   print(myListWithoutSlash[i])
#   row_cells = table.add_row().cells
#   row_cells[2].text = myListWithoutSlash[i]
#   p = row_cells[3].add_paragraph()
#   r = p.add_run()
#   r.add_picture(myList[i],width=Cm(4.0), height=Cm(4))
# document.save('./demo.docx')